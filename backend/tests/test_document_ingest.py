"""Integration — document ingest worker path with mocked LingZhi HTTP."""

from __future__ import annotations

import hashlib
from io import BytesIO
from unittest.mock import patch

import pytest
from docx import Document

from services.domain import jobs as queue
from services.shared.config import S3_BUCKET_DOCUMENTS
from services.shared.db import db_cursor
from services.storage import document_key, get_store
from services.worker.__main__ import handle_document_ingest


@pytest.mark.usefixtures("require_postgres", "require_minio")
def test_document_ingest_offline_ready_without_token(unique_suffix: str, monkeypatch):
    import services.worker.__main__ as worker_mod

    monkeypatch.setattr(worker_mod, "LINGZHI_CLIENT_TOKEN", "")
    monkeypatch.setattr(worker_mod, "LINGZHI_SOURCE_ID", "")

    get_store().ensure_buckets()
    doc = Document()
    doc.add_paragraph("offline ingest unit")
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    sha = hashlib.sha256(data).hexdigest()
    doc_id = f"doc-{unique_suffix}"
    key = document_key("camp-v03", doc_id, sha, "offline.docx")
    get_store().put_bytes(
        S3_BUCKET_DOCUMENTS,
        key,
        data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    with db_cursor() as cur:
        cur.execute(
            """
            INSERT INTO documents
            (id, camp_id, filename, content_type, size_bytes, sha256, object_key, status, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?, 'queued', NOW(), NOW())
            """,
            (
                doc_id,
                "camp-v03",
                "offline.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                len(data),
                sha,
                key,
            ),
        )
    job_id = queue.enqueue_job("document_ingest", {"document_id": doc_id}, camp_id="camp-v03")
    job = queue.get_job(job_id)
    handle_document_ingest(job)
    with db_cursor() as cur:
        cur.execute("SELECT status, error_message FROM documents WHERE id=?", (doc_id,))
        row = cur.fetchone()
    assert row["status"] == "ready"
    assert "LINGZHI" in (row["error_message"] or "")
    assert queue.get_job(job_id)["status"] == "succeeded"


@pytest.mark.usefixtures("require_postgres", "require_minio")
def test_document_ingest_live_publish_mocked(unique_suffix: str, monkeypatch):
    import services.worker.__main__ as worker_mod

    monkeypatch.setattr(worker_mod, "LINGZHI_CLIENT_TOKEN", "test-token")
    monkeypatch.setattr(worker_mod, "LINGZHI_SOURCE_ID", "src-test")
    monkeypatch.setattr(worker_mod, "LINGZHI_BASE_URL", "http://lingzhi.test")

    get_store().ensure_buckets()
    doc = Document()
    doc.add_paragraph("mocked live ingest")
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    sha = hashlib.sha256(data).hexdigest()
    doc_id = f"doc-live-{unique_suffix}"
    key = document_key("camp-v03", doc_id, sha, "live.docx")
    get_store().put_bytes(
        S3_BUCKET_DOCUMENTS,
        key,
        data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    with db_cursor() as cur:
        cur.execute(
            """
            INSERT INTO documents
            (id, camp_id, filename, content_type, size_bytes, sha256, object_key, status, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?, 'queued', NOW(), NOW())
            """,
            (
                doc_id,
                "camp-v03",
                "live.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                len(data),
                sha,
                key,
            ),
        )
    job_id = queue.enqueue_job("document_ingest", {"document_id": doc_id}, camp_id="camp-v03")
    job = queue.get_job(job_id)

    class _Resp:
        def __init__(self, payload, status=200):
            self._payload = payload
            self.status_code = status

        def raise_for_status(self):
            if self.status_code >= 400:
                raise RuntimeError("http error")

        def json(self):
            return self._payload

    class _Client:
        def __init__(self, *a, **k):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        def post(self, url, headers=None, json=None):
            if url.endswith("/check"):
                return _Resp({"results": []})
            if url.endswith("/publish"):
                return _Resp(
                    {
                        "job_id": "lz-job-1",
                        "knowledge_id": "kz-1",
                        "file_id": "fz-1",
                        "status": "queued",
                    }
                )
            raise AssertionError(url)

        def get(self, url, headers=None):
            return _Resp({"status": "completed"})

    with patch("httpx.Client", _Client):
        handle_document_ingest(job)

    with db_cursor() as cur:
        cur.execute(
            "SELECT status, lingzhi_knowledge_id, lingzhi_job_id FROM documents WHERE id=?",
            (doc_id,),
        )
        row = cur.fetchone()
    assert row["status"] == "ready"
    assert row["lingzhi_knowledge_id"] == "kz-1"
    assert row["lingzhi_job_id"] == "lz-job-1"
    assert queue.get_job(job_id)["status"] == "succeeded"
